from __future__ import annotations

import os
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def _ensure_styles(document: Document) -> None:
    styles = document.styles
    if "Heading 1" in styles:
        heading = styles["Heading 1"]
        font = heading.font
        font.name = "Calibri"
        font.size = Pt(20)
        font.bold = True

    for style_name in ("Normal", "BodyText"):
        if style_name in styles:
            style = styles[style_name]
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                font = style.font
                font.name = "Calibri"
                font.size = Pt(11)


def _split_lines(text: str) -> List[str]:
    lines = [line.rstrip() for line in text.splitlines()]
    return [line for line in lines if line.strip()]


def generate_docx_from_resume_text(optimized_resume: str, output_dir: str | os.PathLike[str]) -> str:
    output_path = Path(output_dir).joinpath("optimized_resume.docx")
    document = Document()
    _ensure_styles(document)

    lines = _split_lines(optimized_resume)
    if not lines:
        lines = ["Optimized Resume"]

    name_line = lines[0]
    document.add_heading(name_line, level=1)

    body_lines = lines[1:] or []
    if body_lines:
        for line in body_lines:
            if not line.strip():
                continue
            if line.strip().endswith(":"):
                document.add_heading(line.strip(), level=2)
            else:
                document.add_paragraph(line.strip(), style="BodyText" if "BodyText" in document.styles else "Normal")
    else:
        document.add_paragraph("Summary and details to be added.", style="Normal")

    document.save(str(output_path))
    return str(output_path)

